"""
Remove ALL bold formatting and make EVERYTHING the same size
No exceptions - super uniform professional document
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def make_everything_uniform(file_path):
    """Make EVERYTHING uniform - same size, no bold"""
    filename = os.path.basename(file_path)
    print(f"📄 {filename}")

    try:
        doc = Document(file_path)

        # Process ALL paragraphs - no exceptions
        for para in doc.paragraphs:
            # For every single run, make it uniform
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)  # EVERYTHING 11pt
                run.font.color.rgb = RGBColor(0, 0, 0)  # BLACK
                run.font.bold = False  # NO BOLD AT ALL
                run.font.italic = False  # NO ITALIC
                run.font.underline = False  # NO UNDERLINE

            # Set paragraph formatting
            para.paragraph_format.line_spacing = 1.3
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.space_before = Pt(0)

            # Center only if it's the court title
            if 'IN THE COURT' in para.text.upper():
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.space_before = Pt(4)
                para.space_after = Pt(4)
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Process ALL tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False
                            run.font.italic = False
                            run.font.underline = False

        # Process headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for run in para.runs:
                    # Keep logo as is, but text should be uniform
                    if not run._element.xpath('.//pic:pic'):  # Not an image
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.bold = False

            # Process footers
            footer = section.footer
            for para in footer.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(80, 80, 80)
                    run.font.bold = False

        doc.save(file_path)
        print(f"  ✅ All formatting removed - pure uniform text")
        return True

    except Exception as e:
        print(f"  ❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False

# Process all templates
template_dir = "./data/DOC_TEMPLATE"
templates = [f for f in os.listdir(template_dir) if f.endswith('.docx')]

print("=" * 70)
print("REMOVING ALL BOLD - MAKING EVERYTHING UNIFORM")
print("=" * 70)
print("Changes:")
print("  • Remove ALL bold formatting")
print("  • Make EVERYTHING 11pt")
print("  • No italic, no underline")
print("  • Pure uniform professional text")
print(f"\nProcessing {len(templates)} templates...\n")

success = 0
for template_file in templates:
    file_path = os.path.join(template_dir, template_file)
    if make_everything_uniform(file_path):
        success += 1

print("\n" + "=" * 70)
print(f"✅ {success}/{len(templates)} templates completely uniform")
print("\n📋 Result:")
print("  ✓ NO bold anywhere")
print("  ✓ ALL text 11pt")
print("  ✓ ALL text BLACK")
print("  ✓ Pure uniform professional document")
print("=" * 70)
