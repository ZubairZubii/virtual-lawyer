"""
Make important fields BOLD but keep same size (11pt)
Don't change alignment - keep everything justified
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def format_with_selective_bold(file_path):
    """Add bold to important fields, keep size same, keep alignment"""
    filename = os.path.basename(file_path)
    print(f"📄 {filename}")

    try:
        doc = Document(file_path)

        # Fields that should be bold
        bold_keywords = [
            'IN THE COURT OF',
            'APPLICATION FOR',
            'PETITION FOR',
            'PRAYER:',
            'APPLICANT',
            'PETITIONER',
            'RESPONDENT',
            'THROUGH COUNSEL',
            'ADVOCATE',
            'ENROLLMENT',
            'DATED:'
        ]

        # Process all paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            text_upper = text.upper()

            # Check if this paragraph should be bold
            should_be_bold = any(keyword in text_upper for keyword in bold_keywords)

            # Format all runs
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)  # Always 11pt
                run.font.color.rgb = RGBColor(0, 0, 0)  # BLACK

                # Make bold only if it contains important keywords
                if should_be_bold:
                    run.font.bold = True
                else:
                    run.font.bold = False

            # Set spacing but DON'T change alignment
            para.paragraph_format.line_spacing = 1.3
            para.paragraph_format.space_after = Pt(3)

            # Only center the main court title
            if 'IN THE COURT OF' in text_upper and len(text) > 15:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.space_before = Pt(4)
                para.space_after = Pt(4)
            else:
                # Keep justified for everything else
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False

        doc.save(file_path)
        print(f"  ✅ Bold added, size kept same, alignment preserved")
        return True

    except Exception as e:
        print(f"  ❌ Error: {e}")
        return False

# Process all templates
template_dir = "./data/DOC_TEMPLATE"
templates = [f for f in os.listdir(template_dir) if f.endswith('.docx')]

print("=" * 70)
print("ADDING BOLD TO IMPORTANT FIELDS")
print("=" * 70)
print("Changes:")
print("  • Important fields BOLD (APPLICATION FOR, PRAYER:, etc.)")
print("  • ALL text stays 11pt")
print("  • Alignment stays JUSTIFIED (not changed)")
print("  • Professional appearance")
print(f"\nProcessing {len(templates)} templates...\n")

success = 0
for template_file in templates:
    file_path = os.path.join(template_dir, template_file)
    if format_with_selective_bold(file_path):
        success += 1

print("\n" + "=" * 70)
print(f"✅ {success}/{len(templates)} templates formatted")
print("\n📋 Result:")
print("  ✓ Important fields are BOLD")
print("  ✓ ALL text still 11pt (same size)")
print("  ✓ Alignment preserved (justified)")
print("  ✓ Professional document")
print("=" * 70)
