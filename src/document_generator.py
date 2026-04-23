"""
Document Generation Module
Handles template loading, placeholder filling, and document generation
"""
import os
import json
import re
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime

# Template processing
try:
    from docxtpl import DocxTemplate
    TEMPLATE_AVAILABLE = True
except ImportError:
    TEMPLATE_AVAILABLE = False

# For AI-generated sections
try:
    from multi_layer_pipeline import MultiLayerPipeline
    from multi_source_rag import MultiSourceRAG
    AI_AVAILABLE = True
except ImportError:
    AI_AVAILABLE = False

# Import simplified field configuration
try:
    from src.simplified_field_config import get_simplified_fields, map_simplified_to_original, AUTO_FILL_VALUES
    SIMPLIFIED_CONFIG_AVAILABLE = True
except ImportError:
    try:
        from simplified_field_config import get_simplified_fields, map_simplified_to_original, AUTO_FILL_VALUES
        SIMPLIFIED_CONFIG_AVAILABLE = True
    except ImportError:
        SIMPLIFIED_CONFIG_AVAILABLE = False
        print("Warning: Simplified field config not available")

class DocumentGenerator:
    """
    Document Generation System:
    1. Load template from template library
    2. Extract placeholders
    3. Fill placeholders with user data
    4. Generate AI sections (arguments, case brief, etc.)
    5. Merge template + AI content
    6. Export DOCX/PDF
    """
    
    def __init__(self,
                 template_dir: str = "./data/DOC_TEMPLATE",
                 output_dir: str = "./data/generated_documents",
                 peft_model_path: str = "./models/fine-tuned/golden_model/final_golden_model"):
        """
        Initialize Document Generator
        
        Args:
            template_dir: Directory containing document templates
            output_dir: Directory to save generated documents
            peft_model_path: Path to fine-tuned model for AI generation
        """
        self.template_dir = Path(template_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize AI pipeline if available
        self.ai_pipeline = None
        self.rag = None
        if AI_AVAILABLE:
            try:
                self.ai_pipeline = MultiLayerPipeline(peft_model_path=peft_model_path)
                self.rag = MultiSourceRAG()
            except Exception as e:
                print(f"Warning: AI pipeline not available: {e}")
        
        # Load available templates
        self.templates = self._load_templates()
    
    def _load_templates(self) -> Dict[str, Dict]:
        """Load all available templates from template directory"""
        templates = {}
        
        if not self.template_dir.exists():
            print(f"Warning: Template directory not found: {self.template_dir}")
            return templates
        
        # Search for DOCX files
        for template_file in self.template_dir.rglob("*.docx"):
            # Get category (folder name)
            category = template_file.parent.name if template_file.parent != self.template_dir else "general"
            
            # Get template name
            template_name = template_file.stem
            
            # Extract placeholders and guide
            placeholder_info = self._extract_placeholders(str(template_file))
            placeholders_dict = placeholder_info.get('placeholders', {})
            placeholder_guide = placeholder_info.get('guide', {})
            
            # Get list of placeholder keys
            placeholder_keys = list(placeholders_dict.keys())
            
            templates[f"{category}/{template_name}"] = {
                'file_path': str(template_file),
                'category': category,
                'name': template_name,
                'placeholders': placeholder_keys,
                'placeholder_map': placeholders_dict,  # Maps key -> original name
                'placeholder_guide': placeholder_guide,  # Maps key -> {original_name, description}
                'full_path': str(template_file)
            }
        
        return templates
    
    def _extract_placeholders(self, template_path: str) -> Dict[str, str]:
        """
        Extract placeholders from template
        Supports both formats:
        - {{placeholder_name}} (Jinja2/docxtpl format)
        - {PLACEHOLDER NAME} (single brace format with spaces)
        
        Returns:
            Dict mapping placeholder_key -> original_placeholder_name
        """
        placeholders_dict = {}
        placeholder_guide = {}
        
        try:
            from docx import Document
            doc = Document(template_path)
            # Extract text from all paragraphs
            text = '\n'.join([p.text for p in doc.paragraphs])
            
            # Also check tables for placeholders
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += '\n' + cell.text
            
            # Find all {PLACEHOLDER} patterns (single braces, can have spaces and special chars)
            single_brace_matches = re.findall(r'\{([^}]+)\}', text)
            
            # Find all {{placeholder}} patterns (double braces, Jinja2 format)
            double_brace_matches = re.findall(r'\{\{(\w+)\}\}', text)
            
            # Process single brace placeholders (e.g., {CERTIFICATE REFERENCE NUMBER})
            for match in single_brace_matches:
                # Skip if it's part of a double brace pattern
                if not match.strip().startswith('{') and not match.strip().endswith('}'):
                    original = match.strip()
                    # Convert to key format: "CERTIFICATE REFERENCE NUMBER" -> "certificate_reference_number"
                    key = original.upper().replace(' ', '_').replace('-', '_').replace('/', '_')
                    # Remove special characters but keep underscores
                    key = re.sub(r'[^A-Z0-9_]', '', key)
                    placeholders_dict[key] = original
            
            # Process double brace placeholders (Jinja2 format)
            for match in double_brace_matches:
                key = match.strip()
                placeholders_dict[key] = key
            
            # Try to extract placeholder guide from template
            # Look for "PLACEHOLDER GUIDE" section
            guide_pattern = r'PLACEHOLDER\s+GUIDE[^:]*:(.*?)(?=\n\n|\Z)'
            guide_match = re.search(guide_pattern, text, re.IGNORECASE | re.DOTALL)
            
            if guide_match:
                guide_text = guide_match.group(1)
                # Extract {PLACEHOLDER}: [Description] patterns
                guide_items = re.findall(r'\{([^}]+)\}:\s*\[([^\]]+)\]', guide_text)
                for placeholder_name, description in guide_items:
                    key = placeholder_name.upper().replace(' ', '_').replace('-', '_').replace('/', '_')
                    key = re.sub(r'[^A-Z0-9_]', '', key)
                    placeholder_guide[key] = {
                        'original_name': placeholder_name,
                        'description': description.strip()
                    }
        
        except Exception as e:
            print(f"Error extracting placeholders from {template_path}: {e}")
        
        return {
            'placeholders': placeholders_dict,
            'guide': placeholder_guide
        }
    
    def list_templates(self, category: Optional[str] = None) -> List[Dict]:
        """
        List available templates with SIMPLIFIED placeholder counts

        Args:
            category: Filter by category (optional)

        Returns:
            List of template info dicts with simplified placeholder counts
        """
        templates_list = []

        for template_id, info in self.templates.items():
            if category and info['category'] != category:
                continue

            # Get simplified placeholder count if available
            template_name = info['name']
            placeholder_count = len(info.get('placeholders', []))

            if SIMPLIFIED_CONFIG_AVAILABLE:
                simplified_fields = get_simplified_fields(template_name)
                if simplified_fields:
                    # Use simplified count instead of original
                    placeholder_count = len(simplified_fields)

            # Create template info
            template_info = {
                'id': template_id,
                'name': info['name'],
                'category': info['category'],
                'placeholders': [],  # Don't send full list in summary
                'placeholder_count': placeholder_count  # Use simplified count
            }

            templates_list.append(template_info)

        return templates_list

    def get_template_details_simplified(self, template_id: str) -> Dict:
        """
        Get simplified template details with only essential fields
        Returns user-friendly field names and descriptions

        Args:
            template_id: Template ID

        Returns:
            Dict with simplified fields, descriptions, and metadata
        """
        if template_id not in self.templates:
            raise ValueError(f"Template not found: {template_id}")

        template_info = self.templates[template_id]
        template_name = template_info['name']

        # Try to get simplified fields if available
        simplified_fields = {}
        if SIMPLIFIED_CONFIG_AVAILABLE:
            simplified_fields = get_simplified_fields(template_name)

        # If we have simplified config, use it
        if simplified_fields:
            # Build field list with metadata
            fields_list = []
            placeholder_descriptions = {}
            example_data = {}

            for field_key, field_config in simplified_fields.items():
                fields_list.append(field_key)

                placeholder_descriptions[field_key] = {
                    'original_name': field_config['label'],
                    'description': field_config['description'],
                    'type': field_config['type'],
                    'required': field_config['required']
                }

                # Generate example data
                if field_config['type'] == 'date':
                    example_data[field_key] = "01/01/2024"
                elif field_config['type'] == 'phone':
                    example_data[field_key] = "0300-1234567"
                elif field_config['type'] == 'email':
                    example_data[field_key] = "example@email.com"
                elif field_config['type'] == 'textarea':
                    example_data[field_key] = f"[Enter {field_config['label'].lower()}]"
                else:
                    example_data[field_key] = f"[Enter {field_config['label'].lower()}]"

            return {
                'template_id': template_id,
                'name': template_name,
                'category': template_info.get('category', 'general'),
                'placeholders': fields_list,
                'placeholder_descriptions': placeholder_descriptions,
                'example_data': example_data,
                'total_placeholders': len(fields_list),
                'simplified': True  # Indicate this uses simplified fields
            }

        # Fallback to original complex fields if no simplified config
        return {
            'template_id': template_id,
            'name': template_name,
            'category': template_info.get('category', 'general'),
            'placeholders': template_info.get('placeholders', []),
            'placeholder_map': template_info.get('placeholder_map', {}),
            'placeholder_guide': template_info.get('placeholder_guide', {}),
            'placeholder_descriptions': self._generate_placeholder_descriptions(template_info),
            'example_data': self._generate_example_data(template_info),
            'total_placeholders': len(template_info.get('placeholders', [])),
            'simplified': False  # Original complex fields
        }

    def _generate_placeholder_descriptions(self, template_info: Dict) -> Dict:
        """Generate descriptions for placeholders"""
        descriptions = {}
        placeholder_map = template_info.get('placeholder_map', {})
        placeholder_guide = template_info.get('placeholder_guide', {})

        for key in template_info.get('placeholders', []):
            original_name = placeholder_map.get(key, key)
            guide_info = placeholder_guide.get(key, {})

            # Determine field type from name
            field_type = 'text'
            if any(x in original_name.upper() for x in ['DATE', 'TIME']):
                field_type = 'date'
            elif any(x in original_name.upper() for x in ['EMAIL', 'E-MAIL']):
                field_type = 'email'
            elif any(x in original_name.upper() for x in ['PHONE', 'CONTACT', 'MOBILE']):
                field_type = 'tel'
            elif any(x in original_name.upper() for x in ['ADDRESS', 'DESCRIPTION', 'DETAIL', 'SUMMARY', 'FACTS', 'GROUNDS']):
                field_type = 'textarea'

            descriptions[key] = {
                'original_name': original_name,
                'description': guide_info.get('description', f'Enter {original_name.lower()}'),
                'type': field_type,
                'required': False
            }

        return descriptions

    def _generate_example_data(self, template_info: Dict) -> Dict:
        """Generate example data for placeholders"""
        example_data = {}
        placeholder_map = template_info.get('placeholder_map', {})

        for key in template_info.get('placeholders', []):
            original_name = placeholder_map.get(key, key)
            example_data[key] = f"[{original_name}]"

        return example_data
    
    def generate_ai_section(self, section_type: str, context: Dict) -> str:
        """
        Generate AI-powered sections like:
        - Case brief
        - Legal arguments
        - Grounds for bail
        - Prayer section
        - Case laws
        
        Args:
            section_type: Type of section to generate
            context: Context data (sections, facts, etc.)
        
        Returns:
            Generated text
        """
        if not self.ai_pipeline:
            return ""  # Return empty if AI not available
        
        # Build prompt based on section type
        prompts = {
            'case_brief': f"""Write a concise case brief for a criminal case:
Sections: {context.get('sections', [])}
FIR Number: {context.get('fir_number', 'N/A')}
Police Station: {context.get('police_station', 'N/A')}
Facts: {context.get('facts', 'N/A')}

Write a professional case brief (2-3 paragraphs):""",
            
            'bail_arguments': f"""Write legal arguments for bail application:
Sections: {context.get('sections', [])}
Case Type: {context.get('case_type', 'Criminal')}
Facts: {context.get('facts', 'N/A')}

Write strong legal arguments for bail (3-4 points):""",
            
            'grounds_for_bail': f"""Write grounds for bail application:
Sections: {context.get('sections', [])}
Case Details: {context.get('case_details', 'N/A')}

List 5-7 grounds for bail:""",
            
            'prayer': f"""Write a prayer section for {context.get('document_type', 'petition')}:
Request: {context.get('request', 'Grant bail')}
Sections: {context.get('sections', [])}

Write a professional prayer section:""",
            
            'case_laws': f"""Suggest relevant case laws for:
Sections: {context.get('sections', [])}
Issue: {context.get('issue', 'Bail application')}

List 3-5 relevant case laws with citations:""",
        }
        
        prompt = prompts.get(section_type, f"Generate {section_type} based on: {context}")
        
        try:
            result = self.ai_pipeline.generate_answer(prompt)
            return result.get('answer', '')
        except Exception as e:
            print(f"Error generating AI section: {e}")
            return ""
    
    def fill_template(self,
                     template_id: str,
                     data: Dict,
                     generate_ai_sections: bool = True) -> Dict:
        """
        Fill template with data and generate document
        
        Args:
            template_id: Template ID (e.g., "criminal/pre_arrest_bail")
            data: Data dictionary to fill placeholders
            generate_ai_sections: Whether to generate AI sections
        
        Returns:
            Dict with output_path, status, etc.
        """
        if template_id not in self.templates:
            raise ValueError(f"Template not found: {template_id}")
        
        template_info = self.templates[template_id]
        template_path = template_info['file_path']
        
        if not TEMPLATE_AVAILABLE:
            raise ValueError("Template processing not available. Install python-docxtpl")
        
        # Get placeholder mapping
        placeholder_map = template_info.get('placeholder_map', {})
        template_name = template_info['name']

        # Debug: Print what we received
        print(f"\n=== FILLING TEMPLATE: {template_id} ===")
        print(f"Template name: {template_name}")
        print(f"Received data keys: {list(data.keys())}")

        # Check if we need to map simplified fields to original placeholders
        if SIMPLIFIED_CONFIG_AVAILABLE:
            simplified_config = get_simplified_fields(template_name)
            if simplified_config and any(key in simplified_config for key in data.keys()):
                print("Using simplified field mapping...")
                # User provided simplified field names, map them to original
                expanded_data = map_simplified_to_original(template_name, data)
                print(f"Mapped {len(data)} simplified fields to {len(expanded_data)} original placeholders")
                data = expanded_data

        print(f"Data keys after mapping: {list(data.keys())}")

        # Map user data keys to template placeholder names
        # User provides keys like "certificate_reference_number" or "CERTIFICATE_REFERENCE_NUMBER"
        # Template has "{CERTIFICATE REFERENCE NUMBER}"
        mapped_data = {}
        for user_key, value in data.items():
            if not value or str(value).strip() == "":  # Skip empty values
                continue
                
            # Try to find matching placeholder
            user_key_upper = user_key.upper().replace(' ', '_').replace('-', '_')
            user_key_upper = re.sub(r'[^A-Z0-9_]', '', user_key_upper)
            
            # Check if user_key matches a placeholder key
            if user_key_upper in placeholder_map:
                original_name = placeholder_map[user_key_upper]
                mapped_data[original_name] = str(value).strip()
                print(f"Mapped: {user_key} -> {original_name} = {value}")
            elif user_key in placeholder_map:
                original_name = placeholder_map[user_key]
                mapped_data[original_name] = str(value).strip()
                print(f"Mapped: {user_key} -> {original_name} = {value}")
            else:
                # Try to find by original name directly
                found = False
                for key, orig_name in placeholder_map.items():
                    if orig_name.upper().replace(' ', '_').replace('-', '_') == user_key_upper:
                        mapped_data[orig_name] = str(value).strip()
                        print(f"Mapped (by name): {user_key} -> {orig_name} = {value}")
                        found = True
                        break
                if not found:
                    # Use as-is (might be a direct placeholder name)
                    mapped_data[user_key] = str(value).strip()
                    print(f"Using as-is: {user_key} = {value}")
        
        print(f"Final mapped_data: {mapped_data}")
        
        # Prepare data dictionary for template rendering
        template_data = mapped_data.copy()
        
        # Generate AI sections if requested
        if generate_ai_sections and self.ai_pipeline:
            # Generate case brief
            if 'case_brief' in template_info['placeholders'] and 'case_brief' not in template_data:
                template_data['case_brief'] = self.generate_ai_section('case_brief', data)
            
            # Generate bail arguments
            if 'bail_arguments' in template_info['placeholders'] and 'bail_arguments' not in template_data:
                template_data['bail_arguments'] = self.generate_ai_section('bail_arguments', data)
            
            # Generate grounds
            if 'grounds_for_bail' in template_info['placeholders'] and 'grounds_for_bail' not in template_data:
                template_data['grounds_for_bail'] = self.generate_ai_section('grounds_for_bail', data)
            
            # Generate prayer
            if 'prayer' in template_info['placeholders'] and 'prayer' not in template_data:
                template_data['prayer'] = self.generate_ai_section('prayer', {
                    **data,
                    'document_type': template_info['name']
                })
            
            # Generate case laws
            if 'case_laws' in template_info['placeholders'] and 'case_laws' not in template_data:
                template_data['case_laws'] = self.generate_ai_section('case_laws', data)
        
        # For single-brace placeholders, we need to replace them directly in the document
        # Load the document and replace {PLACEHOLDER} patterns
        use_docxtpl = False
        doc = None
        
        def replace_in_paragraph(paragraph, replacements):
            """Replace all placeholders in a paragraph - handles runs properly"""
            full_text = paragraph.text
            if not full_text:
                return False

            new_text = full_text
            replacements_made = False

            for original_name, value in replacements.items():
                placeholder_pattern = f"{{{original_name}}}"
                if placeholder_pattern in new_text:
                    new_text = new_text.replace(placeholder_pattern, str(value))
                    replacements_made = True
                    print(f"  Replaced {placeholder_pattern} with '{value}'")

            if replacements_made:
                # Preserve paragraph formatting by using the first run's style
                # Clear all existing runs
                first_run_style = None
                if paragraph.runs:
                    first_run_style = paragraph.runs[0].style
                    first_run_font = paragraph.runs[0].font

                # Remove all runs
                for run in paragraph.runs:
                    run.text = ''

                # Set the new text in the first run (or create one if needed)
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    # Restore original formatting
                    if first_run_style:
                        paragraph.runs[0].style = first_run_style
                else:
                    # If no runs exist, create a new one
                    run = paragraph.add_run(new_text)

            return replacements_made
        
        def remove_placeholder_guide_section(doc):
            """Remove the PLACEHOLDER GUIDE section from the document"""
            guide_keywords = [
                "PLACEHOLDER GUIDE",
                "PLACEHOLDER GUIDE FOR USER INPUT",
                "placeholder guide",
                "placeholder guide for user input"
            ]
            
            paragraphs_to_remove = []
            found_guide = False
            guide_start_index = None
            
            # First pass: find where the guide section starts
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip().upper()
                
                # Check if this paragraph starts the guide section
                if any(keyword in para_text for keyword in guide_keywords):
                    found_guide = True
                    guide_start_index = i
                    paragraphs_to_remove.append(i)
                    break
            
            # Second pass: remove all guide entries after the header
            if found_guide and guide_start_index is not None:
                # Remove all paragraphs from guide start to end of document
                # But be smart: stop if we hit a completely blank paragraph followed by non-guide content
                for i in range(guide_start_index + 1, len(doc.paragraphs)):
                    para = doc.paragraphs[i]
                    para_text = para.text.strip()
                    
                    # Stop if we hit a substantial non-guide paragraph (more than 50 chars and doesn't look like guide)
                    if len(para_text) > 50 and not any(
                        '[' in para_text or 
                        ':' in para_text and len(para_text.split(':')) == 2 or
                        para_text.upper().startswith('{') and para_text.upper().endswith('}')
                        for _ in [True]
                    ):
                        # This looks like actual content, stop removing
                        break
                    
                    # Remove guide entries (contain brackets, colons with descriptions, or placeholder patterns)
                    if (para_text and (
                        '[' in para_text or 
                        (':' in para_text and len(para_text.split(':')) == 2) or
                        (para_text.startswith('{') and para_text.endswith('}'))
                    )):
                        paragraphs_to_remove.append(i)
                    elif para_text.strip() == '':
                        # Empty line, might be separator, remove it
                        paragraphs_to_remove.append(i)
            
            # Remove paragraphs in reverse order to maintain indices
            for i in reversed(paragraphs_to_remove):
                try:
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
                except:
                    pass  # Paragraph might already be removed
            
            # Also check tables for guide sections
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip().upper()
                        if any(keyword in cell_text for keyword in guide_keywords):
                            # Clear the entire cell
                            for para in cell.paragraphs:
                                para.clear()
                            break
        
        try:
            from docx import Document
            doc = Document(template_path)
            
            print(f"\nReplacing placeholders in document...")
            print(f"Template data to replace: {template_data}")
            replacements_count = 0
            
            # Replace placeholders in paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text:
                    if replace_in_paragraph(paragraph, template_data):
                        replacements_count += 1
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text:
                                if replace_in_paragraph(para, template_data):
                                    replacements_count += 1
            
            print(f"Made {replacements_count} replacements in document")

            # Remove paragraphs with unreplaced placeholders
            print("Removing paragraphs with unreplaced placeholders...")
            paragraphs_to_remove = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                # Check if paragraph contains unreplaced placeholders like {SOMETHING}
                if text and '{' in text and '}' in text:
                    # Check if it's an actual unreplaced placeholder
                    if re.search(r'\{[A-Z_0-9]+\}', text):
                        print(f"  Removing paragraph with unreplaced placeholder: {text[:80]}...")
                        paragraphs_to_remove.append(i)

            # Remove in reverse order
            for i in reversed(paragraphs_to_remove):
                try:
                    p = doc.paragraphs[i]._element
                    p.getparent().remove(p)
                except:
                    pass

            # Remove placeholder guide section
            print("Removing placeholder guide section...")
            remove_placeholder_guide_section(doc)
            print("Placeholder guide section removed")
            
            # Successfully used direct replacement
            use_docxtpl = False
        except Exception as e:
            print(f"Error with direct replacement, trying DocxTemplate: {e}")
            import traceback
            traceback.print_exc()
            use_docxtpl = True
        
        if use_docxtpl:
            # Fallback to DocxTemplate for Jinja2 format {{placeholder}}
            doc = DocxTemplate(template_path)
            # Fill missing placeholders with defaults
            for placeholder_key in template_info['placeholders']:
                original_name = placeholder_map.get(placeholder_key, placeholder_key)
                if original_name not in template_data and placeholder_key not in template_data:
                    template_data[placeholder_key] = f"[{original_name}]"
            doc.render(template_data)
            
        # Generate output filename first
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{template_info['name']}_{timestamp}.docx"
        output_path = self.output_dir / output_filename
        
        if use_docxtpl:
            # Fallback to DocxTemplate for Jinja2 format {{placeholder}}
            doc_template = DocxTemplate(template_path)
            # Fill missing placeholders with defaults
            for placeholder_key in template_info['placeholders']:
                original_name = placeholder_map.get(placeholder_key, placeholder_key)
                if original_name not in template_data and placeholder_key not in template_data:
                    template_data[placeholder_key] = f"[{original_name}]"
            doc_template.render(template_data)
            doc_template.save(str(output_path))
            
            # Remove guide section from saved document
            try:
                from docx import Document
                doc_clean = Document(str(output_path))
                remove_placeholder_guide_section(doc_clean)
                doc_clean.save(str(output_path))
                print("Placeholder guide section removed from DocxTemplate document")
            except Exception as e:
                print(f"Warning: Could not remove guide section: {e}")
        else:
            # Save document (already has guide removed)
            doc.save(str(output_path))
        
        # PDF conversion disabled for now - only DOCX support
        pdf_path = None
        pdf_filename = None
        
        # Return only filenames (not full paths) for frontend
        return {
            'output_path': str(output_path),
            'output_filename': output_filename,
            'pdf_path': str(pdf_path) if pdf_path and pdf_path.exists() else None,
            'pdf_filename': pdf_filename if pdf_path and pdf_path.exists() else None,  # Just filename, not path
            'template_id': template_id,
            'template_name': template_info['name'],
            'status': 'success',
            'placeholders_filled': len([p for p in template_info['placeholders'] if p in template_data]),
            'total_placeholders': len(template_info['placeholders'])
        }
    
    def suggest_document_type(self, facts: Dict) -> List[str]:
        """
        Suggest document types based on extracted facts
        
        Args:
            facts: Extracted facts from document analysis
        
        Returns:
            List of suggested template IDs
        """
        suggestions = []
        
        sections = facts.get('sections', [])
        fir_number = facts.get('fir_number')
        
        # Logic to suggest documents based on sections
        if any('302' in str(s) for s in sections):  # Murder
            suggestions.append('criminal/post_arrest_bail')
            suggestions.append('criminal/quash_fir')
        
        if any('380' in str(s) or '457' in str(s) for s in sections):  # Theft
            suggestions.append('criminal/pre_arrest_bail')
            suggestions.append('criminal/497_bail')
        
        if any('376' in str(s) for s in sections):  # Rape
            suggestions.append('criminal/post_arrest_bail')
        
        if fir_number:
            suggestions.append('criminal/22A_application')
            suggestions.append('criminal/22B_application')
        
        # Remove duplicates and return
        return list(set(suggestions))
    
    def validate_data(self, template_id: str, data: Dict) -> Dict:
        """
        Validate data before generating document

        Returns:
            Dict with validation status, missing_fields, warnings
        """
        if template_id not in self.templates:
            return {
                'valid': False,
                'error': f"Template not found: {template_id}"
            }

        template_info = self.templates[template_id]
        required_placeholders = template_info['placeholders']

        missing = []
        warnings = []

        for placeholder in required_placeholders:
            if placeholder not in data or not data[placeholder]:
                missing.append(placeholder)

        # Check for critical fields (only warn, don't block generation)
        critical_fields = ['accused_name', 'fir_number', 'sections']
        for field in critical_fields:
            if field in required_placeholders and (field not in data or not data[field]):
                warnings.append(f"Critical field missing: {field}")

        # Allow document generation even with missing fields
        # Only fail if template is invalid (which we already checked above)
        return {
            'valid': True,  # Always valid as long as template exists
            'missing_fields': missing,
            'warnings': warnings,
            'filled_fields': len([p for p in required_placeholders if p in data and data[p]]),
            'has_missing_fields': len(missing) > 0
        }





















