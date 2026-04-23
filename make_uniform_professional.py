"""
Make ALL templates uniform and professional
- Same size throughout (11pt)
- Minimal bold (only main title)
- All text black
- Super professional appearance
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def make_all_text_uniform(para, is_main_title=False):
    """Make all text uniform - same size, mostly not bold"""

    # Clear all formatting first
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)  # BLACK
        run.font.bold = False  # Remove bold by default
        run.font.size = Pt(11)  # Standard size

    # Only the very first main title gets special treatment
    if is_main_title:
        for run in para.runs:
            run.font.size = Pt(12)  # Slightly bigger
            run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.space_before = Pt(6)
        para.space_after = Pt(6)
    else:
        # Everything else is normal
        para.paragraph_format.line_spacing = 1.3
        para.paragraph_format.space_after = Pt(3)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def make_document_uniform(file_path):
    """Make entire document uniform"""
    filename = os.path.basename(file_path)
    print(f"📄 {filename}")

    try:
        doc = Document(file_path)

        main_title_found = False

        # Process all paragraphs
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()

            # Only the FIRST "IN THE COURT" gets to be title
            if not main_title_found and 'IN THE COURT' in text and len(text) > 15:
                make_all_text_uniform(para, is_main_title=True)
                main_title_found = True
            else:
                # Everything else is uniform normal text
                make_all_text_uniform(para, is_main_title=False)

        # Process tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(11)
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.font.bold = False

        doc.save(file_path)
        print(f"  ✅ All text now uniform")
        return True

    except Exception as e:
        print(f"  ❌ Error: {e}")
        return False

# Process all templates
template_dir = "./data/DOC_TEMPLATE"
templates = [f for f in os.listdir(template_dir) if f.endswith('.docx')]

print("=" * 70)
print("MAKING ALL TEXT UNIFORM AND PROFESSIONAL")
print("=" * 70)
print("Changes:")
print("  • ALL text 11pt (only main title 12pt)")
print("  • Remove excessive BOLD")
print("  • Everything BLACK")
print("  • Uniform professional appearance")
print(f"\nProcessing {len(templates)} templates...\n")

success = 0
for template_file in templates:
    file_path = os.path.join(template_dir, template_file)
    if make_document_uniform(file_path):
        success += 1

print("\n" + "=" * 70)
print(f"✅ {success}/{len(templates)} templates made uniform")
print("\n📋 Result:")
print("  ✓ All text same size (11pt)")
print("  ✓ Only main title is bold (12pt)")
print("  ✓ Everything else normal weight")
print("  ✓ All BLACK text")
print("  ✓ Super professional appearance")
print("=" * 70)
